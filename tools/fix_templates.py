"""
Fix template formatting issues and add {{NOTICE_HEADER}} placeholder.

Task 1 — Commercial template formatting:
  1. "To:" and tenant name on same line → "To: {{TENANT_NAMES}}"
  2. Property address left-aligned flush with "To:" line
  3. DUE DATE / AMOUNT DUE section → proper 2-column Word table
  4. TOTAL AMOUNT DUE → same line, centered

Task 2 — Add {{NOTICE_HEADER}} to both templates:
  Replace hardcoded "THREE (3) DAY NOTICE TO PAY OR QUIT" with {{NOTICE_HEADER}}
"""

from docx import Document
from docx.shared import Inches, Pt, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'templates')


def find_paragraph_index(doc, text_match, start=0):
    """Find index of paragraph containing exact text."""
    for i, p in enumerate(doc.paragraphs):
        if i >= start and text_match in p.text:
            return i
    return None


def copy_run_format(src_run, dst_run):
    """Copy font formatting from source run to destination run."""
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb


def remove_paragraph(paragraph):
    """Remove a paragraph element from the document."""
    p = paragraph._element
    p.getparent().remove(p)


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'none'))
        element.set(qn('w:sz'), val.get('sz', '0'))
        element.set(qn('w:space'), val.get('space', '0'))
        element.set(qn('w:color'), val.get('color', 'auto'))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def make_invisible_borders(cell):
    """Set all borders to invisible."""
    border = {'val': 'none', 'sz': '0', 'space': '0', 'color': 'auto'}
    set_cell_border(cell, top=border, bottom=border, start=border, end=border)


def fix_commercial_template(filepath):
    """Fix formatting issues in commercial template."""
    doc = Document(filepath)
    paragraphs = doc.paragraphs

    # === TASK 2: Replace title with {{NOTICE_HEADER}} ===
    title_idx = find_paragraph_index(doc, 'THREE (3) DAY NOTICE TO PAY OR QUIT')
    if title_idx is not None:
        title_p = paragraphs[title_idx]
        # Preserve the style/formatting of the paragraph
        style = title_p.style
        alignment = title_p.alignment
        pf = title_p.paragraph_format

        # Get formatting from first non-empty run
        ref_run = None
        for r in title_p.runs:
            if r.text.strip():
                ref_run = r
                break

        # Clear all runs
        for r in title_p.runs:
            r.text = ''
        # Set text on first run or add new one
        if title_p.runs:
            title_p.runs[0].text = '{{NOTICE_HEADER}}'
        else:
            run = title_p.add_run('{{NOTICE_HEADER}}')
            if ref_run:
                copy_run_format(ref_run, run)
        # Remove extra empty runs
        p_elem = title_p._element
        runs = p_elem.findall(qn('w:r'))
        for r in runs[1:]:
            if r.text is None or r.text.strip() == '':
                # Check if the run has actual text
                t_elem = r.find(qn('w:t'))
                if t_elem is None or not t_elem.text or not t_elem.text.strip():
                    p_elem.remove(r)

    # === TASK 1.1: Fix "To:" line — merge with tenant names ===
    to_idx = find_paragraph_index(doc, 'To:')
    tenant_idx = find_paragraph_index(doc, '{{TENANT_NAMES}}')
    if to_idx is not None and tenant_idx is not None:
        to_p = paragraphs[to_idx]
        tenant_p = paragraphs[tenant_idx]

        # Get formatting from existing "To:" run
        to_run = to_p.runs[0] if to_p.runs else None

        # Change "To:" to "To: {{TENANT_NAMES}}"
        if to_p.runs:
            to_p.runs[0].text = 'To: {{TENANT_NAMES}}'
        else:
            run = to_p.add_run('To: {{TENANT_NAMES}}')

        # Remove the separate {{TENANT_NAMES}} paragraph
        remove_paragraph(tenant_p)

    # Re-index paragraphs after removal
    paragraphs = doc.paragraphs

    # === TASK 1.2: Fix property address alignment ===
    # Find "To:" paragraph's left indent and apply same to address paragraphs
    to_idx = find_paragraph_index(doc, 'To:')
    if to_idx is not None:
        to_indent = paragraphs[to_idx].paragraph_format.left_indent or Emu(225425)
        addr_street_idx = find_paragraph_index(doc, '{{PROPERTY_ADDRESS_STREET}}')
        addr_city_idx = find_paragraph_index(doc, '{{PROPERTY_ADDRESS_CITY}}')
        for idx in [addr_street_idx, addr_city_idx]:
            if idx is not None:
                paragraphs[idx].paragraph_format.left_indent = to_indent

    # === TASK 1.3: Convert DUE DATE / AMOUNT DUE to a 2-column table ===
    due_date_idx = find_paragraph_index(doc, 'DUE DATE')
    amount_due_idx = find_paragraph_index(doc, 'AMOUNT DUE')
    rent_due_idx = find_paragraph_index(doc, '{{RENT_DUE_DATE}}')

    if due_date_idx is not None and amount_due_idx is not None and rent_due_idx is not None:
        # Get the font size from the header runs
        header_size = Pt(11.5)  # 146050 EMU ≈ 11.5pt
        for r in paragraphs[due_date_idx].runs:
            if r.font.size:
                header_size = r.font.size
                break

        # Insert table BEFORE the DUE DATE paragraph
        # We need to insert the table element in the document body
        due_date_elem = paragraphs[due_date_idx]._element
        parent = due_date_elem.getparent()

        # Create the table
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(2.5)

        # Header row
        h_due = table.cell(0, 0).paragraphs[0]
        h_due.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h_due.add_run('DUE DATE')
        run.bold = True
        run.font.size = header_size

        h_amt = table.cell(0, 1).paragraphs[0]
        h_amt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h_amt.add_run('AMOUNT DUE')
        run.bold = True
        run.font.size = header_size

        # Data row
        d_due = table.cell(1, 0).paragraphs[0]
        d_due.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = d_due.add_run('{{RENT_DUE_DATE}}')
        run.font.size = header_size

        d_amt = table.cell(1, 1).paragraphs[0]
        d_amt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = d_amt.add_run('{{AMOUNT_DUE}}')
        run.font.size = header_size

        # Make borders invisible (clean look)
        for row in table.rows:
            for cell in row.cells:
                make_invisible_borders(cell)

        # Move the table element to before the DUE DATE paragraph
        tbl_elem = table._tbl
        # Remove from current position (it was added at end)
        tbl_elem.getparent().remove(tbl_elem)
        # Insert before DUE DATE paragraph
        parent.insert(parent.index(due_date_elem), tbl_elem)

        # Remove the old DUE DATE, AMOUNT DUE, and data paragraphs
        # Re-index after table insertion
        paragraphs = doc.paragraphs
        # Find them again
        for p in list(doc.paragraphs):
            if p.text in ('DUE DATE', 'AMOUNT DUE') or '{{RENT_DUE_DATE}}' in p.text:
                remove_paragraph(p)

    # Re-index
    paragraphs = doc.paragraphs

    # === TASK 1.4: Fix TOTAL AMOUNT DUE — centered, no tab ===
    total_idx = find_paragraph_index(doc, 'TOTAL AMOUNT DUE:')
    if total_idx is not None:
        total_p = paragraphs[total_idx]
        # Replace tab with space, center the paragraph
        new_text = 'TOTAL AMOUNT DUE: {{TOTAL_AMOUNT_DUE}}'

        # Get first run's formatting
        ref_run = None
        for r in total_p.runs:
            if r.text.strip():
                ref_run = r
                break

        # Clear existing runs
        p_elem = total_p._element
        for r in list(p_elem.findall(qn('w:r'))):
            p_elem.remove(r)

        # Add single run with corrected text
        run = total_p.add_run(new_text)
        if ref_run:
            copy_run_format(ref_run, run)

        # Center the paragraph and remove left indent
        total_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_p.paragraph_format.left_indent = None

    doc.save(filepath)
    print(f'Fixed: {os.path.basename(filepath)}')


def fix_residential_template(filepath):
    """Add {{NOTICE_HEADER}} and apply same formatting fixes to residential template."""
    doc = Document(filepath)
    paragraphs = doc.paragraphs

    # === TASK 2: Replace title with {{NOTICE_HEADER}} ===
    title_idx = find_paragraph_index(doc, 'THREE (3) DAY NOTICE TO PAY OR QUIT')
    if title_idx is not None:
        title_p = paragraphs[title_idx]
        ref_run = None
        for r in title_p.runs:
            if r.text.strip():
                ref_run = r
                break

        for r in title_p.runs:
            r.text = ''
        if title_p.runs:
            title_p.runs[0].text = '{{NOTICE_HEADER}}'
        else:
            run = title_p.add_run('{{NOTICE_HEADER}}')
            if ref_run:
                copy_run_format(ref_run, run)

        p_elem = title_p._element
        runs = p_elem.findall(qn('w:r'))
        for r in runs[1:]:
            t_elem = r.find(qn('w:t'))
            if t_elem is None or not t_elem.text or not t_elem.text.strip():
                p_elem.remove(r)

    # === Fix "To:" line — merge with tenant names ===
    to_idx = find_paragraph_index(doc, 'To:')
    tenant_idx = find_paragraph_index(doc, '{{TENANT_NAMES}}')
    if to_idx is not None and tenant_idx is not None:
        to_p = paragraphs[to_idx]
        tenant_p = paragraphs[tenant_idx]
        if to_p.runs:
            to_p.runs[0].text = 'To: {{TENANT_NAMES}}'
        else:
            to_p.add_run('To: {{TENANT_NAMES}}')
        remove_paragraph(tenant_p)

    # Re-index
    paragraphs = doc.paragraphs

    # === Fix property address alignment ===
    to_idx = find_paragraph_index(doc, 'To:')
    if to_idx is not None:
        to_indent = paragraphs[to_idx].paragraph_format.left_indent
        if to_indent:
            addr_street_idx = find_paragraph_index(doc, '{{PROPERTY_ADDRESS_STREET}}')
            addr_city_idx = find_paragraph_index(doc, '{{PROPERTY_ADDRESS_CITY}}')
            for idx in [addr_street_idx, addr_city_idx]:
                if idx is not None:
                    paragraphs[idx].paragraph_format.left_indent = to_indent

    # === Convert DUE DATE / AMOUNT DUE to table ===
    due_date_idx = find_paragraph_index(doc, 'DUE DATE')
    amount_due_idx = find_paragraph_index(doc, 'AMOUNT DUE')
    rent_due_idx = find_paragraph_index(doc, '{{RENT_DUE_DATE}}')

    if due_date_idx is not None and amount_due_idx is not None and rent_due_idx is not None:
        header_size = Pt(11.5)
        for r in paragraphs[due_date_idx].runs:
            if r.font.size:
                header_size = r.font.size
                break

        due_date_elem = paragraphs[due_date_idx]._element
        parent = due_date_elem.getparent()

        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(2.5)

        h_due = table.cell(0, 0).paragraphs[0]
        h_due.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h_due.add_run('DUE DATE')
        run.bold = True
        run.font.size = header_size

        h_amt = table.cell(0, 1).paragraphs[0]
        h_amt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h_amt.add_run('AMOUNT DUE')
        run.bold = True
        run.font.size = header_size

        d_due = table.cell(1, 0).paragraphs[0]
        d_due.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = d_due.add_run('{{RENT_DUE_DATE}}')
        run.font.size = header_size

        d_amt = table.cell(1, 1).paragraphs[0]
        d_amt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = d_amt.add_run('{{AMOUNT_DUE}}')
        run.font.size = header_size

        for row in table.rows:
            for cell in row.cells:
                make_invisible_borders(cell)

        tbl_elem = table._tbl
        tbl_elem.getparent().remove(tbl_elem)
        parent.insert(parent.index(due_date_elem), tbl_elem)

        paragraphs = doc.paragraphs
        for p in list(doc.paragraphs):
            if p.text in ('DUE DATE', 'AMOUNT DUE') or '{{RENT_DUE_DATE}}' in p.text:
                remove_paragraph(p)

    # Re-index
    paragraphs = doc.paragraphs

    # === Fix TOTAL AMOUNT DUE ===
    total_idx = find_paragraph_index(doc, 'TOTAL AMOUNT DUE:')
    if total_idx is not None:
        total_p = paragraphs[total_idx]
        new_text = 'TOTAL AMOUNT DUE: {{TOTAL_AMOUNT_DUE}}'
        ref_run = None
        for r in total_p.runs:
            if r.text.strip():
                ref_run = r
                break

        p_elem = total_p._element
        for r in list(p_elem.findall(qn('w:r'))):
            p_elem.remove(r)

        run = total_p.add_run(new_text)
        if ref_run:
            copy_run_format(ref_run, run)

        total_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_p.paragraph_format.left_indent = None

    doc.save(filepath)
    print(f'Fixed: {os.path.basename(filepath)}')


if __name__ == '__main__':
    commercial = os.path.join(
        TEMPLATES_DIR,
        '3-Day Notice - BLUEPRINT - commercial - NEW BRANDING_071125.docx'
    )
    residential = os.path.join(
        TEMPLATES_DIR,
        '3-Day Notice - BLUEPRINT - residential - NEW BRANDING_071125_v2.docx'
    )

    fix_commercial_template(commercial)
    fix_residential_template(residential)
    print('Done. Both templates updated.')
