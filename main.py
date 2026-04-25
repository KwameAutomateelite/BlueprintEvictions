print("STARTUP: importing stdlib...", flush=True)
import json
import logging
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional

print("STARTUP: importing httpx...", flush=True)
import httpx
print("STARTUP: importing python-docx...", flush=True)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt
print("STARTUP: importing pypdf...", flush=True)
from pypdf import PdfReader, PdfWriter
print("STARTUP: importing dropbox-sign...", flush=True)
from dropbox_sign import ApiClient, ApiException, Configuration, apis, models
print("STARTUP: importing fastapi...", flush=True)
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import FileResponse, PlainTextResponse, Response
from pydantic import BaseModel

print("STARTUP: all imports OK", flush=True)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Blueprint Evictions - Dropbox Sign Service", version="2.1.0")

# --- Config ---

print("STARTUP: reading env vars...", flush=True)
print(f"STARTUP: DROPBOX_SIGN_API_KEY set = {'DROPBOX_SIGN_API_KEY' in os.environ}", flush=True)
print(f"STARTUP: AIRTABLE_API_KEY set = {'AIRTABLE_API_KEY' in os.environ}", flush=True)
print(f"STARTUP: PORT = {os.environ.get('PORT', 'NOT SET')}", flush=True)
DROPBOX_SIGN_API_KEY = os.environ["DROPBOX_SIGN_API_KEY"]
AIRTABLE_API_KEY = os.environ["AIRTABLE_API_KEY"]
AIRTABLE_BASE_ID = os.environ.get("AIRTABLE_BASE_ID", "appumajkmYLcMryFd")
AIRTABLE_TABLE_ID = os.environ.get("AIRTABLE_TABLE_ID", "tblZbUGz8OTvFNh9i")

configuration = Configuration(username=DROPBOX_SIGN_API_KEY)

TEMPLATES_DIR = Path(__file__).parent / "templates"
ATTACHMENTS_DIR = Path(__file__).parent / "attachments"

TEMPLATE_MAP = {
    "3-Day Pay or Quit": "3-Day Notice - BLUEPRINT - commercial - NEW BRANDING_071125.docx",
    "3-Day Notice to Pay Rent or Quit": "3-Day Notice - BLUEPRINT - commercial - NEW BRANDING_071125.docx",
    "3 Day Pay or Quit": "3-Day Notice - BLUEPRINT - commercial - NEW BRANDING_071125.docx",
    "3-Day Perform or Quit": "3day_perform_quit_TEMPLATE.docx",
    "3-Day Quit": "3day_quit_notice_TEMPLATE.docx",
    "TPO Warning": "tpo_warning_TEMPLATE.docx",
    "TPA Warning": "tpa_warning_TEMPLATE.docx",
}

NOTICE_TEMPLATES = {
    "residential": "3-Day Notice - BLUEPRINT - residential - NEW BRANDING_071125_v2.docx",
    "commercial": "3-Day Notice - BLUEPRINT - commercial - NEW BRANDING_071125.docx",
}


# --- Models ---


class SendSignatureRequest(BaseModel):
    signer_email: str
    signer_name: str
    document_name: str
    record_id: str
    notice_type: str
    case_name: str
    fields: dict
    file_url: Optional[str] = None
    attachments_required: List[str] = []


class SendSignatureResponse(BaseModel):
    signature_request_id: str
    status: str
    message: str


class AmountDue(BaseModel):
    due_date: str
    amount: str


class GenerateNoticeRequest(BaseModel):
    notice_type: str  # "residential" or "commercial"
    day_count: int  # 3, 5, 10, 15, 30
    tenant_names: str
    property_address: str = ""
    property_address_street: str = ""  # Optional: if provided, used directly instead of splitting property_address
    property_address_city: str = ""    # Optional: if provided, used directly instead of splitting property_address
    county: str
    amounts_due: List[AmountDue]
    total_amount_due: str
    service_date: str = ""
    payment_address: str = ""
    landlord_name: str
    landlord_phone: str = ""
    landlord_address: str = ""
    landlord_city: str = ""
    landlord_state: str = ""
    landlord_zip: str = ""
    notice_date: str = ""
    case_name: str = ""
    notice_header: str = ""  # Override for notice header text; defaults based on day_count
    is_section_8: bool = False
    is_san_jose_tpo: bool = False
    is_mountain_view_csfra: bool = False
    attachments_required: List[str] = []


DAY_COUNT_WORDS = {
    3: ("three", "THREE"),
    5: ("five", "FIVE"),
    10: ("ten", "TEN"),
    15: ("fifteen", "FIFTEEN"),
    30: ("thirty", "THIRTY"),
}

# Blueprint Evictions service-area cities (defensive trailer stripper for street blobs
# where city/state/zip got concatenated into the street field, e.g. Airtable data like
# "123 Main St Santa Cruz CA 98598"). Sorted longest-first at match time so
# "East Palo Alto" matches before "Palo Alto".
KNOWN_CITIES = [
    "San Jose", "Santa Clara", "Sunnyvale", "Mountain View", "Cupertino", "Milpitas",
    "Los Gatos", "Saratoga", "Campbell", "Gilroy", "Morgan Hill", "Palo Alto",
    "Los Altos", "Los Altos Hills", "Monte Sereno", "Alum Rock",
    "East Palo Alto", "Redwood City", "San Mateo", "Daly City", "South San Francisco",
    "Burlingame", "Menlo Park", "Foster City", "San Carlos", "Belmont", "Millbrae",
    "Atherton", "Brisbane", "Colma", "Half Moon Bay", "Hillsborough", "Pacifica",
    "Portola Valley", "San Bruno", "Woodside",
    "Oakland", "Berkeley", "Fremont", "Hayward", "Pleasanton", "Dublin", "Livermore",
    "Union City", "San Leandro", "Alameda", "Newark", "Emeryville", "Albany",
    "Santa Cruz", "Scotts Valley", "Capitola", "Watsonville", "Aptos",
    "Walnut Creek", "Concord", "Richmond", "Antioch", "Pittsburg", "Martinez",
    "San Ramon", "Danville", "Lafayette", "Orinda",
    "San Francisco",
    "Hollister", "San Juan Bautista",
]


def _strip_city_state_zip_trailer(street_blob: str) -> str:
    """Strip a '<City> CA <ZIP>' trailer from a street-address blob.

    Airtable's 'Tenant's Property Address' sometimes contains the full concatenated
    address (e.g. '123 Main St Santa Cruz CA 98598') with no comma delimiter.
    When the workflow also sends a separate city/state/zip line, that full blob ends up
    in {{PROPERTY_ADDRESS_STREET}} causing a duplicate render on the notice.

    This is defensive: n8n should also strip this before sending. If neither runs,
    the trailer remains and you'll see duplication — but never dropped street data.
    """
    if not street_blob:
        return street_blob
    s = street_blob.strip()
    for city in sorted(KNOWN_CITIES, key=len, reverse=True):
        pattern = re.compile(
            r"\s+" + re.escape(city) + r"\s*,?\s*CA\s+\d{5}(?:-\d{4})?\s*$",
            re.IGNORECASE,
        )
        m = pattern.search(s)
        if m:
            return s[: m.start()].rstrip(" ,")
    fallback = re.compile(r"\s+CA\s+\d{5}(?:-\d{4})?\s*$")
    m = fallback.search(s)
    if m:
        return s[: m.start()].rstrip(" ,")
    return s


def _parse_landlord_address(full: str) -> tuple:
    """Split a landlord-address blob into (street_line1, 'City, CA ZIP').

    Handles the smashed-blob case Airtable leaks (e.g. 'Santa CruzCA12345' with no
    delimiters) so the owner block and housing-provider block render two clean rows
    instead of a single run of mashed-together text.

    Order of attempts:
      1. KNOWN_CITIES anchored match (handles 'Santa CruzCA12345' with no separators)
      2. Generic '<stuff> CA <ZIP>' regex (relaxed whitespace)
      3. Comma split ('456 Main St, Santa Cruz, CA 12345')
      4. Newline split ('456 Main St\\nSanta Cruz, CA 12345')
      5. Fallback: return (full, '')

    Returns a tuple of plain strings; never None.
    """
    if not full:
        return ("", "")
    s = full.strip()
    normalized = re.sub(r"\s+", " ", s.replace("\n", " ").replace("\r", " ")).strip()

    for city in sorted(KNOWN_CITIES, key=len, reverse=True):
        pattern = re.compile(
            r"^(.*?)\s*" + re.escape(city) + r"\s*,?\s*(?:CA|California)\s*(\d{5}(?:-\d{4})?)\s*$",
            re.IGNORECASE,
        )
        m = pattern.match(normalized)
        if m:
            line1 = m.group(1).strip(" ,")
            csz = f"{city}, CA {m.group(2)}"
            return (line1, csz)

    generic = re.match(
        r"^(.*?)\s*(?:CA|California)\s+(\d{5}(?:-\d{4})?)\s*$",
        normalized,
        re.IGNORECASE,
    )
    if generic:
        before = generic.group(1).strip(" ,")
        return (before, f"CA {generic.group(2)}")

    if "," in s:
        parts = s.split(",", 1)
        return (parts[0].strip(), parts[1].strip())

    if "\n" in s:
        parts = s.split("\n", 1)
        return (parts[0].strip(), parts[1].strip())

    return (s, "")


# --- Helpers ---


def fill_template(notice_type: str, fields: dict) -> str:
    """Fill a Word template with field values and return the path to the filled .docx."""
    template_name = TEMPLATE_MAP.get(notice_type)
    if not template_name:
        raise ValueError(
            f"Unknown notice_type '{notice_type}'. "
            f"Valid types: {', '.join(TEMPLATE_MAP.keys())}"
        )

    template_path = TEMPLATES_DIR / template_name
    logger.info(f"DEBUG template_path: {template_path}")
    logger.info(f"DEBUG template_path.exists(): {template_path.exists()}")
    logger.info(f"DEBUG TEMPLATES_DIR contents: {list(TEMPLATES_DIR.iterdir()) if TEMPLATES_DIR.exists() else 'DIR NOT FOUND'}")
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    logger.info(f"DEBUG template file size: {os.path.getsize(str(template_path))} bytes")

    doc = Document(str(template_path))
    logger.info(f"DEBUG paragraphs count: {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs[:5]):
        logger.info(f"DEBUG para[{i}]: {p.text[:120]!r}")

    # Replace {{PLACEHOLDER}} tags in paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, fields)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, fields)

    # Replace in headers/footers
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer is not None:
                for paragraph in header_footer.paragraphs:
                    _replace_in_paragraph(paragraph, fields)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc.save(tmp.name)
    saved_size = os.path.getsize(tmp.name)
    logger.info(f"DEBUG saved filled docx: {tmp.name} size={saved_size}")

    # Validate the saved docx is not corrupt
    if saved_size == 0:
        raise RuntimeError(f"fill_template produced empty file: {tmp.name}")
    with open(tmp.name, "rb") as f:
        header = f.read(20)
    logger.info(f"DEBUG docx header hex: {header.hex()}")
    # Valid docx (ZIP) starts with PK (50 4b)
    if not header.startswith(b"PK"):
        raise RuntimeError(
            f"fill_template produced invalid docx (not a ZIP). "
            f"Size={saved_size}, header={header.hex()}"
        )
    # Re-open with python-docx to verify
    try:
        test_doc = Document(tmp.name)
        logger.info(f"DEBUG docx validation OK: {len(test_doc.paragraphs)} paragraphs")
    except Exception as e:
        raise RuntimeError(
            f"fill_template produced corrupt docx: {e}. Size={saved_size}"
        )

    return tmp.name


def _replace_in_paragraph(paragraph, fields: dict) -> None:
    """Replace {{KEY}} placeholders in a paragraph while preserving formatting.

    Handles two cases:
    1. Placeholder fits entirely within one run → replace in-place, formatting preserved.
    2. Placeholder spans multiple runs (Word splits text unpredictably) → merge only
       the affected runs into the first one, replace there, clear the rest.
    """
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(run.text for run in runs)
    if "{{" not in full_text:
        return

    # Build list of placeholders to replace
    replacements = {}
    for key, value in fields.items():
        placeholder = "{{" + key + "}}"
        if placeholder in full_text:
            replacements[placeholder] = str(value) if value else ""

    if not replacements:
        return

    # Pass 1: Replace placeholders that fit within a single run
    for run in runs:
        for placeholder, value in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    # Check if any placeholders still remain (they span multiple runs)
    full_text_after = "".join(run.text for run in runs)
    remaining = {p: v for p, v in replacements.items() if p in full_text_after}

    if not remaining:
        return

    # Pass 2: Handle cross-run placeholders by finding the run span
    for placeholder, value in remaining.items():
        while True:
            # Rebuild positions each iteration since runs change
            texts = [run.text for run in runs]
            combined = "".join(texts)
            pos = combined.find(placeholder)
            if pos < 0:
                break

            # Find which runs the placeholder spans
            char_count = 0
            start_run = end_run = None
            for i, t in enumerate(texts):
                run_start = char_count
                run_end = char_count + len(t)
                if start_run is None and pos < run_end:
                    start_run = i
                    offset_in_start = pos - run_start
                if start_run is not None and pos + len(placeholder) <= run_end:
                    end_run = i
                    offset_in_end = pos + len(placeholder) - run_start
                    break
                char_count = run_end

            if start_run is None or end_run is None:
                break

            # Merge the spanned text into the start run, replace, clear others
            merged = "".join(texts[start_run : end_run + 1])
            runs[start_run].text = merged.replace(placeholder, value, 1)
            for j in range(start_run + 1, end_run + 1):
                runs[j].text = ""


def convert_docx_to_pdf(docx_path: str) -> str:
    """Convert a .docx file to PDF using LibreOffice."""
    logger.info(f"DEBUG convert_docx_to_pdf input: {docx_path} size={os.path.getsize(docx_path)}")
    pdf_path = docx_path.rsplit(".", 1)[0] + ".pdf"
    output_dir = os.path.dirname(docx_path)

    lo_env = os.environ.copy()
    lo_env["HOME"] = "/tmp"

    cmd = [
        "libreoffice", "--headless", "--norestore", "--nofirststartwizard",
        "--convert-to", "pdf", "--outdir", output_dir, docx_path,
    ]
    logger.info(f"DEBUG LibreOffice cmd: {' '.join(cmd)}")

    result = subprocess.run(
        cmd, capture_output=True, text=True, timeout=120, env=lo_env,
    )
    logger.info(f"DEBUG LibreOffice rc={result.returncode} stdout={result.stdout} stderr={result.stderr}")

    if not os.path.exists(pdf_path):
        dir_contents = os.listdir(output_dir) if os.path.isdir(output_dir) else "DIR NOT FOUND"
        raise RuntimeError(
            f"PDF conversion failed (rc={result.returncode}). "
            f"stdout: {result.stdout}. stderr: {result.stderr}. "
            f"Dir contents: {dir_contents}"
        )

    logger.info(f"DEBUG PDF created: {pdf_path} size={os.path.getsize(pdf_path)}")
    return pdf_path


def generate_notice_pdf(notice_type: str, fields: dict) -> str:
    """Fill a Word template and convert to PDF. Returns the PDF temp file path."""
    docx_path = fill_template(notice_type, fields)
    try:
        pdf_path = convert_docx_to_pdf(docx_path)
        return pdf_path
    finally:
        try:
            os.unlink(docx_path)
        except OSError:
            pass


async def download_file(url: str) -> str:
    """Download a file from URL to a temp path and return the path.

    Detects DOCX vs PDF from magic bytes + Content-Type so the Approve path
    can pull the LIVE DOCX from OneDrive (preserving manual edits made in
    Word) instead of regenerating from template. DOCX files are then converted
    to PDF downstream before being sent to Dropbox Sign.
    """
    async with httpx.AsyncClient(follow_redirects=True, timeout=60) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        content = resp.content
        ctype = resp.headers.get("content-type", "").lower()

    # DOCX is a ZIP — starts with "PK". PDF starts with "%PDF".
    is_docx = (
        content[:2] == b"PK"
        or "wordprocessingml" in ctype
        or "vnd.openxmlformats" in ctype
    )
    is_pdf = content[:4] == b"%PDF" or "application/pdf" in ctype

    if is_docx and not is_pdf:
        suffix = ".docx"
    else:
        suffix = ".pdf"

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(content)
    tmp.close()
    logger.info(f"download_file: saved {tmp.name} size={len(content)} ctype={ctype}")
    return tmp.name


async def update_airtable_status(record_id: str, status: str) -> None:
    """Update a record's Status field in Airtable."""
    url = f"https://api.airtable.com/v0/{AIRTABLE_BASE_ID}/{AIRTABLE_TABLE_ID}/{record_id}"
    headers = {
        "Authorization": f"Bearer {AIRTABLE_API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {"fields": {"Status": status}}

    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.patch(url, headers=headers, json=payload)
        if resp.status_code != 200:
            logger.error(
                f"Airtable update failed for {record_id}: {resp.status_code} {resp.text}"
            )
            raise HTTPException(
                status_code=502, detail=f"Airtable update failed: {resp.text}"
            )
        logger.info(f"Airtable record {record_id} updated to Status='{status}'")


def merge_attachment_pdfs(notice_pdf_path: str, attachment_filenames: List[str]) -> str:
    """Merge static attachment PDFs onto the end of the notice PDF.

    Args:
        notice_pdf_path: Path to the generated notice PDF.
        attachment_filenames: List of PDF filenames in ATTACHMENTS_DIR to append.

    Returns:
        Path to the merged PDF (new temp file). Original is not modified.
    """
    if not attachment_filenames:
        return notice_pdf_path

    writer = PdfWriter()

    # Add the notice PDF pages
    notice_reader = PdfReader(notice_pdf_path)
    for page in notice_reader.pages:
        writer.add_page(page)

    # Append each attachment PDF in order
    for filename in attachment_filenames:
        attachment_path = ATTACHMENTS_DIR / filename
        if not attachment_path.exists():
            logger.warning(f"Attachment not found, skipping: {attachment_path}")
            continue
        attachment_reader = PdfReader(str(attachment_path))
        for page in attachment_reader.pages:
            writer.add_page(page)
        logger.info(f"Merged attachment: {filename} ({len(attachment_reader.pages)} pages)")

    # Write merged PDF
    merged_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    merged_tmp.close()
    writer.write(merged_tmp.name)
    merged_size = os.path.getsize(merged_tmp.name)
    logger.info(f"Merged PDF: {merged_tmp.name} size={merged_size} ({len(writer.pages)} total pages)")

    return merged_tmp.name


# --- Endpoints ---


@app.get("/health")
async def health():
    return {"status": "ok", "service": "dropbox-sign-service"}


@app.get("/debug-template")
async def debug_template():
    """Debug endpoint: run template pipeline and return diagnostics."""
    info = {}
    info["TEMPLATES_DIR"] = str(TEMPLATES_DIR)
    info["TEMPLATES_DIR_exists"] = TEMPLATES_DIR.exists()
    info["TEMPLATES_DIR_contents"] = [f.name for f in TEMPLATES_DIR.iterdir()] if TEMPLATES_DIR.exists() else []

    template_name = NOTICE_TEMPLATES["commercial"]
    template_path = TEMPLATES_DIR / template_name
    info["template_path"] = str(template_path)
    info["template_exists"] = template_path.exists()
    info["template_size"] = os.path.getsize(str(template_path)) if template_path.exists() else 0

    logo_path = TEMPLATES_DIR / "blueprint_logo.jpg"
    info["logo_path"] = str(logo_path)
    info["logo_exists"] = logo_path.exists()
    info["logo_size"] = os.path.getsize(str(logo_path)) if logo_path.exists() else 0

    if template_path.exists():
        doc = Document(str(template_path))
        info["paragraph_count"] = len(doc.paragraphs)
        info["first_5_paragraphs"] = [p.text[:150] for p in doc.paragraphs[:5]]
        info["table_count"] = len(doc.tables)

        # Test fill + convert
        try:
            fields = {"TENANT_NAMES": "DEBUG_TENANT", "COUNTY": "DEBUG_COUNTY"}
            docx_path = fill_template("3-Day Pay or Quit", fields)
            info["filled_docx_size"] = os.path.getsize(docx_path)

            pdf_path = convert_docx_to_pdf(docx_path)
            info["pdf_size"] = os.path.getsize(pdf_path)
            info["pdf_path"] = pdf_path

            # Read first 200 bytes of PDF to confirm it's a real PDF
            with open(pdf_path, "rb") as f:
                header = f.read(200)
            info["pdf_header"] = header[:50].decode("latin-1")

            os.unlink(docx_path)
            os.unlink(pdf_path)
        except Exception as e:
            info["error"] = str(e)

    return info


class TestPdfRequest(BaseModel):
    notice_type: str
    fields: dict


@app.post("/test-pdf")
async def test_pdf(req: TestPdfRequest):
    """Generate a filled PDF from template and return it directly (no Dropbox Sign)."""
    try:
        pdf_path = generate_notice_pdf(notice_type=req.notice_type, fields=req.fields)
    except Exception as e:
        logger.error(f"test-pdf failed: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        filename=f"{req.notice_type.replace(' ', '_')}.pdf",
    )


@app.post("/send-signature", response_model=SendSignatureResponse)
async def send_signature(req: SendSignatureRequest):
    """Fill a branded Word template (or download from URL) and send to Dropbox Sign."""
    logger.info(
        f"Sending signature request: {req.document_name} to {req.signer_email}"
    )

    # Get PDF: either download from URL or generate from template
    logger.info(f"DEBUG file_url={req.file_url!r} notice_type={req.notice_type!r}")
    logger.info(f"DEBUG fields keys={list(req.fields.keys())}")
    try:
        if req.file_url:
            logger.info("DEBUG: Taking file_url download branch")
            file_path = await download_file(req.file_url)
        else:
            logger.info("DEBUG: Taking template generation branch")
            file_path = generate_notice_pdf(
                notice_type=req.notice_type,
                fields=req.fields,
            )
    except Exception as e:
        logger.error(f"Failed to prepare PDF: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to prepare PDF: {str(e)}")
    logger.info(f"DEBUG final file_path={file_path} size={os.path.getsize(file_path)}")

    # Normalize to PDF. When file_url points at a live OneDrive DOCX (the Approve
    # path after AM manually edits in Word), download_file saves it with a .docx
    # suffix. Convert to PDF here so Dropbox Sign always receives PDF.
    if file_path.endswith(".docx"):
        docx_path = file_path
        file_path = convert_docx_to_pdf(docx_path)
        try:
            os.unlink(docx_path)
        except OSError:
            pass

    # Merge attachment PDFs if required
    original_file_path = file_path
    if req.attachments_required:
        logger.info(f"Merging {len(req.attachments_required)} attachments: {req.attachments_required}")
        notice_pdf = file_path
        file_path = merge_attachment_pdfs(notice_pdf, req.attachments_required)
        if file_path != notice_pdf:
            try:
                os.unlink(notice_pdf)
            except OSError:
                pass
        logger.info(f"Final merged PDF: {file_path} size={os.path.getsize(file_path)}")

    try:
        with ApiClient(configuration) as api_client:
            signature_request_api = apis.SignatureRequestApi(api_client)

            signer = models.SubSignatureRequestSigner(
                email_address=req.signer_email,
                name=req.signer_name,
                order=0,
            )

            signing_options = models.SubSigningOptions(
                draw=True,
                type=True,
                upload=True,
                phone=False,
                default_type="draw",
            )

            data = models.SignatureRequestSendRequest(
                title=req.document_name,
                subject=f"Please sign: {req.document_name}",
                message=(
                    f"Hi {req.signer_name},\n\n"
                    f"Please review and sign the attached document: {req.document_name}.\n\n"
                    "Thank you,\nBlueprint Evictions LLC"
                ),
                signers=[signer],
                files=[open(file_path, "rb")],
                metadata={"record_id": req.record_id},
                signing_options=signing_options,
                test_mode=os.environ.get("DROPBOX_SIGN_TEST_MODE", "0") == "1",
            )

            result = signature_request_api.signature_request_send(data)
            sig_request = result.signature_request

            logger.info(
                f"Signature request created: {sig_request.signature_request_id}"
            )

            return SendSignatureResponse(
                signature_request_id=sig_request.signature_request_id,
                status="sent",
                message=f"Signature request sent to {req.signer_email}",
            )

    except ApiException as e:
        logger.error(f"Dropbox Sign API error: {e.status} {e.body}")
        raise HTTPException(
            status_code=502,
            detail=f"Dropbox Sign API error: {e.body}",
        )
    finally:
        # Clean up temp file
        try:
            os.unlink(file_path)
        except OSError:
            pass


@app.post("/signature-callback")
async def signature_callback(request: Request):
    """Webhook endpoint that Dropbox Sign calls when a document is signed.

    Dropbox Sign sends event data as form-encoded with a 'json' field.
    On event 'signature_request_all_signed', update Airtable Status to 'Signed (A)'.
    """
    form = await request.form()

    # Dropbox Sign sends the payload in a 'json' form field
    json_str = form.get("json")
    if not json_str:
        logger.warning("Callback received with no 'json' field")
        return {"status": "ignored", "reason": "no json payload"}

    payload = json.loads(json_str)
    event = payload.get("event", {})
    event_type = event.get("event_type")
    event_time = event.get("event_time")

    logger.info(f"Dropbox Sign callback: event_type={event_type} time={event_time}")

    # Respond to the callback test (Dropbox Sign sends this to verify the endpoint)
    if event_type == "callback_test":
        logger.info("Responding to callback_test with 'Hello API Event Received'")
        return PlainTextResponse("Hello API Event Received")

    # Only act on all-signed events
    if event_type != "signature_request_all_signed":
        logger.info(f"Ignoring event type: {event_type}")
        return {"status": "ignored", "event_type": event_type}

    # Extract metadata with record_id
    signature_request = payload.get("signature_request", {})
    metadata = signature_request.get("metadata", {})
    record_id = metadata.get("record_id")
    sig_request_id = signature_request.get("signature_request_id", "unknown")

    if not record_id:
        logger.error(
            f"No record_id in metadata for signature request {sig_request_id}"
        )
        return {"status": "error", "reason": "no record_id in metadata"}

    logger.info(
        f"All signed for request {sig_request_id}, updating record {record_id}"
    )

    # Update Airtable
    await update_airtable_status(record_id, "Signed (A)")

    return {
        "status": "processed",
        "event_type": event_type,
        "record_id": record_id,
        "signature_request_id": sig_request_id,
    }


def _replace_day_count_in_paragraph(paragraph, day_count: int) -> None:
    """Replace day count references in a paragraph, handling split runs.

    Handles both single-run matches (e.g. 'three (3)') and split-run
    matches where 'THREE', ' ', '(3)' are in separate runs.
    """
    if day_count == 3:
        return  # Template already says "three (3)", nothing to replace

    lower_word, upper_word = DAY_COUNT_WORDS.get(day_count, (str(day_count), str(day_count)))
    runs = paragraph.runs
    if not runs:
        return

    # Try single-run replacement first
    for run in runs:
        run.text = run.text.replace("THREE (3)", f"{upper_word} ({day_count})")
        run.text = run.text.replace("three (3)", f"{lower_word} ({day_count})")
        run.text = run.text.replace("3 (three)", f"{day_count} ({lower_word})")

    # Handle split runs: replace individual tokens
    for run in runs:
        if run.text == "THREE":
            run.text = upper_word
        elif run.text == "(3)":
            run.text = f"({day_count})"
        elif run.text == "three":
            run.text = lower_word
        # Also handle inline occurrences like "within three" or "the 3"
        elif "three" in run.text and "three" != lower_word:
            run.text = run.text.replace("three", lower_word)
        elif " 3 " in run.text:
            run.text = run.text.replace(" 3 ", f" {day_count} ")


def _fill_cell_with_values(cell, values: list, template_paragraph) -> None:
    """Fill a table cell with multiple values, one paragraph per value.

    Preserves paragraph alignment and run font formatting from the template paragraph.
    """
    from copy import deepcopy
    from docx.oxml.ns import qn

    # Get formatting from the template paragraph
    template_alignment = template_paragraph.paragraph_format.alignment
    template_run = template_paragraph.runs[0] if template_paragraph.runs else None

    # Set the first value in the existing paragraph (preserves original formatting)
    if template_paragraph.runs:
        template_paragraph.runs[0].text = values[0]
        for run in template_paragraph.runs[1:]:
            run.text = ""
    else:
        template_paragraph.text = values[0]

    # Add additional paragraphs for remaining values, copying formatting.
    # Strip w:spacing from cloned pPr so 3+ amount rows don't stack ~200 twips
    # before-spacing each — which pushes the body long enough to force the
    # California-notices block onto a third page with an orphan page 2 behind it.
    for val in values[1:]:
        new_para = deepcopy(template_paragraph._element)
        cloned_pPr = new_para.find(qn("w:pPr"))
        if cloned_pPr is not None:
            sp = cloned_pPr.find(qn("w:spacing"))
            if sp is not None:
                cloned_pPr.remove(sp)
        # Clear text in the cloned paragraph, then set new value
        for r in new_para.findall(qn("w:r")):
            new_para.remove(r)
        # Create a new run with the value
        new_run_elem = deepcopy(template_run._element) if template_run else None
        if new_run_elem is not None:
            # Clear existing text nodes and set new text
            for t in new_run_elem.findall(qn("w:t")):
                new_run_elem.remove(t)
            t_elem = new_run_elem.makeelement(qn("w:t"), {})
            t_elem.text = val
            # Preserve spaces
            t_elem.set(qn("xml:space"), "preserve")
            new_run_elem.append(t_elem)
            new_para.append(new_run_elem)
        else:
            r_elem = new_para.makeelement(qn("w:r"), {})
            t_elem = r_elem.makeelement(qn("w:t"), {})
            t_elem.text = val
            t_elem.set(qn("xml:space"), "preserve")
            r_elem.append(t_elem)
            new_para.append(r_elem)
        # Insert the new paragraph after the current last paragraph in the cell
        cell._element.append(new_para)


def _fill_amounts_paragraph(doc, amounts_due: list) -> None:
    """Replace {{RENT_DUE_DATE}} / {{AMOUNT_DUE}} placeholders with actual amounts.

    Handles two template layouts:
    1. Tab-separated paragraphs: {{RENT_DUE_DATE}}\t{{AMOUNT_DUE}}
       Uses a right-aligned tab stop so dollar amounts line up.
    2. Word table cells: separate cells for {{RENT_DUE_DATE}} and {{AMOUNT_DUE}}
       Creates one paragraph per amount row, preserving alignment and font formatting.
    """
    if not amounts_due:
        return

    # --- Layout 1: Tab-separated paragraphs (legacy) ---
    for para in doc.paragraphs:
        full_text = "".join(run.text for run in para.runs)
        if "{{RENT_DUE_DATE}}" in full_text and "{{AMOUNT_DUE}}" in full_text:
            # Add a right-aligned tab stop at 6.5 inches for amount column
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

            # Build formatted lines — single tab to hit the right-aligned stop
            lines = []
            for a in amounts_due:
                lines.append(f"{a['due_date']}\t{a['amount']}")
            replacement = "\n".join(lines)
            if para.runs:
                para.runs[0].text = replacement
                for run in para.runs[1:]:
                    run.text = ""
            return  # Done — legacy layout found

    # --- Layout 2: Word table cells ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if "{{RENT_DUE_DATE}}" in cell_text or "{{AMOUNT_DUE}}" in cell_text:
                    # Found the amounts table — fill cells with proper paragraphs
                    dates = [a["due_date"] for a in amounts_due]
                    amounts = [a["amount"] for a in amounts_due]

                    for r in table.rows:
                        for c in r.cells:
                            for para in c.paragraphs:
                                full = "".join(run.text for run in para.runs)
                                if "{{RENT_DUE_DATE}}" in full:
                                    _fill_cell_with_values(c, dates, para)
                                elif "{{AMOUNT_DUE}}" in full:
                                    _fill_cell_with_values(c, amounts, para)
                    return  # Done — table layout found


def _get_paragraph_left_indent_twips(paragraph) -> int:
    """Return the effective left indent (in twips/dxa) of a paragraph, or 0 if none."""
    from docx.oxml.ns import qn

    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return 0
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return 0
    left = ind.get(qn("w:left")) or ind.get(qn("w:start"))
    try:
        return int(left) if left is not None else 0
    except ValueError:
        return 0


# Landlord-block indent: ~3.75" left indent (5400 twips) so the block visually
# anchors on the RIGHT side of the page near the signature line — per AM's
# reference image where Kwame Amankwah / 122 First Addy Dr / Santa Cruz CA 12345
# sit right of center. Body width is 10800 twips on this template (asymmetric
# pgMar), so a 5400 indent puts the block's left edge at the horizontal midpoint;
# typical landlord-line text widths (~30-40 chars) then trail off to the right
# margin, giving a right-anchored visual.
LANDLORD_BLOCK_LEFT_TWIPS = 5400

# When AMOUNTS_DUE has at least this many rows, force a hard page break before
# the California notices heading. Below that threshold the soft compaction
# logic in _compact_empty_paragraphs_before_heading already keeps the doc to
# 2 pages without extra whitespace on page 1.
FORCE_PAGE_BREAK_ROW_THRESHOLD = 6


def _ensure_pPr(paragraph):
    """Return the paragraph's <w:pPr>, creating it (as the first child) if absent."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._p.insert(0, pPr)
    return pPr


def _set_paragraph_spacing_explicit(paragraph, before=None, after=None, line=None) -> None:
    """Force explicit w:spacing values. Pass None to leave a value alone."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = _ensure_pPr(paragraph)
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    if before is not None:
        sp.set(qn("w:before"), str(before))
    if after is not None:
        sp.set(qn("w:after"), str(after))
    if line is not None:
        sp.set(qn("w:line"), str(line))


def _set_paragraph_indent(paragraph, left=None, right=None, first_line=None) -> None:
    """Force explicit w:ind values."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = _ensure_pPr(paragraph)
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if left is not None:
        ind.set(qn("w:left"), str(left))
    if right is not None:
        ind.set(qn("w:right"), str(right))
    if first_line is not None:
        ind.set(qn("w:firstLine"), str(first_line))


def _set_paragraph_alignment(paragraph, value: str) -> None:
    """Force w:jc to the given value ('left', 'center', 'right', 'both')."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = _ensure_pPr(paragraph)
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), value)


def _force_left_align_doc(doc) -> None:
    """Force every body and table-cell paragraph to w:jc=left.

    Per AM's call: 'everything left-justified, nothing centered' for body text.
    The Blueprint logo lives in section headers (not iterated here) and the
    rent table itself stays centered on the page (table-level w:jc); only the
    text inside paragraphs is normalized.
    """

    for paragraph in doc.paragraphs:
        _set_paragraph_alignment(paragraph, "left")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _set_paragraph_alignment(paragraph, "left")


_ALL_OTHER_OCCUPANTS_RE = re.compile(
    r"\b(?:and\s+)?all\s+other\s+occupants\b",
    re.IGNORECASE,
)


def _normalize_all_other_occupants(text: str) -> str:
    """Force 'and All Other Occupants' (Title Case A/O/O) regardless of input case."""
    if not text:
        return text

    def _fix(match: "re.Match[str]") -> str:
        original = match.group(0)
        prefix = "and " if original.lower().startswith("and ") else ""
        return f"{prefix}All Other Occupants"

    return _ALL_OTHER_OCCUPANTS_RE.sub(_fix, text)


def _enforce_all_other_occupants_in_doc(doc) -> None:
    """Defense-in-depth: rewrite any 'all other occupants' run text to title case.

    Belt-and-suspenders against an upstream LLM that emits the wrong casing.
    Run-level rewrite (not paragraph-level) so we don't disturb other formatting.
    """
    def _fix_paragraph(paragraph) -> None:
        for run in paragraph.runs:
            new_text = _normalize_all_other_occupants(run.text)
            if new_text != run.text:
                run.text = new_text

    for paragraph in doc.paragraphs:
        _fix_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _fix_paragraph(paragraph)


TENANT_BLOCK_LEFT_TWIPS = 720  # 0.5" — matches AM-approved reference image indent.


def _fix_tenant_address_block(doc) -> None:
    """Group street/city/COUNTY OF as a single visual unit on page 1.

    Template ships with three separate paragraphs for the property address
    (PROPERTY_ADDRESS_STREET, PROPERTY_ADDRESS_CITY, COUNTY OF {{COUNTY}}),
    plus an empty spacer paragraph between the city line and COUNTY OF, and
    a different left-indent on the COUNTY line. After placeholder substitution:

      1. add visible spacing AFTER the 'situated at:' line,
      2. zero before/after spacing on street and city lines so they render
         as one tight block,
      3. drop the empty spacer between city and COUNTY OF,
      4. apply a uniform 0.5" indent (TENANT_BLOCK_LEFT_TWIPS) to all three
         lines so the block sits visibly indented from body text per the
         AM-approved reference image, then add visible spacing AFTER the
         COUNTY line (paragraph break before '[hereinafter called the
         "Premises"]').
    """
    body = doc.element.body
    paras = list(doc.paragraphs)

    situated_idx = None
    county_idx = None
    for i, p in enumerate(paras):
        text = p.text
        if situated_idx is None and "situated" in text and "at:" in text:
            situated_idx = i
        elif p.text.strip().upper().startswith("COUNTY OF"):
            county_idx = i
            break

    if situated_idx is None or county_idx is None:
        logger.info("tenant-address-fix: skipping (could not locate landmarks)")
        return

    _set_paragraph_spacing_explicit(paras[situated_idx], after=240)
    _set_paragraph_spacing_explicit(paras[county_idx], before=0, after=240)
    _set_paragraph_indent(paras[county_idx], left=TENANT_BLOCK_LEFT_TWIPS)
    _set_paragraph_alignment(paras[county_idx], "left")

    for j in range(situated_idx + 1, county_idx):
        text = paras[j].text.strip()
        if not text:
            body.remove(paras[j]._element)
        else:
            _set_paragraph_spacing_explicit(paras[j], before=0, after=0)
            _set_paragraph_indent(paras[j], left=TENANT_BLOCK_LEFT_TWIPS)
            _set_paragraph_alignment(paras[j], "left")


# Rent table sizing: AM-approved reference shows the table occupying ~65-70%
# of body width, centered as a block on the page (DUE DATE col starts ~25%
# from left margin, AMOUNT DUE col right edge sits ~10% from right margin).
# 7200 twips (~5") split into two 3600-twip columns, table-level w:jc=center.
RENT_TABLE_BODY_WIDTH_TWIPS = 7200
RENT_TABLE_COL_WIDTH_TWIPS = RENT_TABLE_BODY_WIDTH_TWIPS // 2  # 3600 each


def _set_rent_table_column_widths(doc) -> None:
    """Force the DUE DATE / AMOUNT DUE table to span body width with two equal columns.

    Sets table-level fixed layout (so widths stick rather than auto-fitting to
    content), full body-width w:tblW, fresh 2-column w:tblGrid, matching per-cell
    w:tcW on every cell, AND overrides the template's leaked w:jc=center on the
    table to w:jc=left so a sub-body-width table doesn't drift inward when the
    section's body actually has spare room.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for table in doc.tables:
        if len(table.columns) != 2 or not table.rows:
            continue
        first_row_text = " ".join(c.text for c in table.rows[0].cells).upper()
        if "DUE DATE" not in first_row_text or "AMOUNT DUE" not in first_row_text:
            continue

        tbl = table._element

        # Ensure tblPr exists as the first child of <w:tbl>.
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Fixed layout — widths are honored verbatim instead of auto-sizing.
        tbl_layout = tblPr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tblPr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")

        # AM-approved reference: table sits CENTERED as a block on the page
        # (sub-body-width — 7200 of 10800 twips). Table-level w:jc=center is
        # the correct setting here; do not flip to left.
        tbl_jc = tblPr.find(qn("w:jc"))
        if tbl_jc is None:
            tbl_jc = OxmlElement("w:jc")
            tblPr.append(tbl_jc)
        tbl_jc.set(qn("w:val"), "center")

        # Drop any tblInd that would shift the table inward from body-left.
        tbl_ind = tblPr.find(qn("w:tblInd"))
        if tbl_ind is not None:
            tbl_ind.set(qn("w:w"), "0")
            tbl_ind.set(qn("w:type"), "dxa")

        # Total table width = body width (dxa = twips).
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), str(RENT_TABLE_BODY_WIDTH_TWIPS))
        tblW.set(qn("w:type"), "dxa")

        # Replace tblGrid with two equal-width columns.
        old_grid = tbl.find(qn("w:tblGrid"))
        if old_grid is not None:
            tbl.remove(old_grid)
        new_grid = OxmlElement("w:tblGrid")
        for _ in range(2):
            gridCol = OxmlElement("w:gridCol")
            gridCol.set(qn("w:w"), str(RENT_TABLE_COL_WIDTH_TWIPS))
            new_grid.append(gridCol)
        tblPr.addnext(new_grid)

        # Set tcW on every cell so per-cell widths line up with the grid.
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    tc.insert(0, tcPr)
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"), str(RENT_TABLE_COL_WIDTH_TWIPS))
                tcW.set(qn("w:type"), "dxa")
        return


def _align_rent_table_columns(doc) -> None:
    """Set DUE DATE column paragraphs CENTER, AMOUNT DUE column paragraphs RIGHT.

    AM-approved reference: dates ("January 1, 2026", etc.) center within the
    DUE DATE column under the centered "DUE DATE" header. Dollar amounts
    right-align under the AMOUNT DUE header so digits stack on a shared right
    edge (e.g. $7,000.00 / $25,000.00). The TOTAL row's "TOTAL AMOUNT DUE:"
    label also lands in col 0 and so picks up CENTER — matching the target
    where the label sits just left of center.

    MUST run AFTER `_force_left_align_doc` (which flattens every cell to LEFT)
    AND AFTER `_add_total_row_to_amounts_table` (so the new TOTAL row also
    gets the column convention). This is intentionally the LAST alignment
    pass on the rent table.
    """
    for table in doc.tables:
        if len(table.columns) != 2:
            continue
        first_row_text = " ".join(c.text for c in table.rows[0].cells).upper()
        if "DUE DATE" not in first_row_text or "AMOUNT DUE" not in first_row_text:
            continue
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 1:
                for paragraph in cells[0].paragraphs:
                    _set_paragraph_alignment(paragraph, "center")
            if len(cells) >= 2:
                for paragraph in cells[1].paragraphs:
                    _set_paragraph_alignment(paragraph, "right")
        return


def _add_total_row_to_amounts_table(doc, total_amount_due: str) -> bool:
    """Move 'TOTAL AMOUNT DUE: $X' into the rent table as a final 2-cell row.

    Layout per AM-approved reference: 'TOTAL AMOUNT DUE:' sits LEFT-aligned in
    column 1 (under DUE DATE column), and the dollar value sits RIGHT-aligned in
    column 2 (under AMOUNT DUE column, anchored at the right margin like a
    spreadsheet total). Both cells get w:before=360 spacing (~0.25", ~one line
    height) so there's a visible blank-line gap above the TOTAL row.

    Returns True if the table-row swap happened, False if the rent table or
    the standalone TOTAL paragraph couldn't be located.
    """
    from docx.oxml.ns import qn

    target_table = None
    for table in doc.tables:
        if len(table.columns) != 2 or not table.rows:
            continue
        first_row_text = " ".join(c.text for c in table.rows[0].cells).upper()
        if "DUE DATE" in first_row_text and "AMOUNT DUE" in first_row_text:
            target_table = table
            break

    if target_table is None:
        return False

    # Visible blank spacer row — gives ~1 line of vertical space between the
    # last itemized row and the TOTAL row, matching AM's reference. More
    # reliable than tuning w:before spacing (which kept rendering too tight).
    spacer_row = target_table.add_row()
    for spacer_cell in spacer_row.cells:
        for sp in spacer_cell.paragraphs:
            for r in list(sp.runs):
                r.text = ""
            _set_paragraph_spacing_explicit(sp, before=0, after=0)

    new_row = target_table.add_row()
    label_cell, value_cell = new_row.cells[0], new_row.cells[1]

    label_para = label_cell.paragraphs[0]
    for r in list(label_para.runs):
        r.text = ""
    run = label_para.add_run("TOTAL AMOUNT DUE:")
    run.bold = True
    run.font.size = Pt(11.5)
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing_explicit(label_para, before=0, after=0)

    value_para = value_cell.paragraphs[0]
    for r in list(value_para.runs):
        r.text = ""
    vrun = value_para.add_run(total_amount_due or "")
    vrun.bold = True
    vrun.font.size = Pt(11.5)
    value_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing_explicit(value_para, before=0, after=0)

    from docx.oxml import OxmlElement

    body = doc.element.body
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if not text.upper().startswith("TOTAL AMOUNT DUE"):
            continue
        total_p = para._element

        # Both templates park empty BodyText paragraphs between the rent table
        # and the standalone TOTAL paragraph (residential: 1; commercial: 5+)
        # to push TOTAL down the page. With TOTAL relocated into the table,
        # those empties become a 1–2 line dead gap above "You are further
        # notified". Walk back from TOTAL and drop each preceding empty <w:p>
        # until we hit non-empty content or the table.
        prev = total_p.getprevious()
        while prev is not None and prev.tag == qn("w:p"):
            prev_text = "".join(t.text or "" for t in prev.findall(".//" + qn("w:t"))).strip()
            if prev_text:
                break
            to_remove = prev
            prev = prev.getprevious()
            body.remove(to_remove)

        # Normalize the "You are further notified" paragraph's w:before to one
        # line (240 twips at 11pt). Residential ships at 257, commercial at
        # 150 — pinning to 240 gives a consistent single-line gap below the
        # in-table TOTAL row across both templates.
        nxt = total_p.getnext()
        if nxt is not None and nxt.tag == qn("w:p"):
            pPr = nxt.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                nxt.insert(0, pPr)
            sp = pPr.find(qn("w:spacing"))
            if sp is None:
                sp = OxmlElement("w:spacing")
                pPr.append(sp)
            sp.set(qn("w:before"), "240")

        body.remove(total_p)
        return True
    return True


def _force_page_break_before(doc, heading_substring: str) -> bool:
    """Set <w:pageBreakBefore/> on the first body paragraph matching the heading.

    Used to guarantee 'NOTICES FROM THE STATE OF CALIFORNIA' starts on page 2
    regardless of how many AMOUNTS_DUE rows are rendered above. Without the
    forced break, 12-row notices push California (and everything below it)
    to page 3 — AM's 'three pages of a notice' failure mode.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for para in doc.paragraphs:
        if heading_substring not in para.text:
            continue
        pPr = _ensure_pPr(para)
        pbb = pPr.find(qn("w:pageBreakBefore"))
        if pbb is None:
            pbb = OxmlElement("w:pageBreakBefore")
            pPr.append(pbb)
        return True
    return False


def _insert_paragraphs_before(
    insert_before_element,
    lines: list,
    ref_paragraph,
    zero_spacing: bool = False,
    left_indent: Optional[int] = None,
) -> None:
    """Insert N plain paragraphs (one per string in `lines`) before the given XML element.

    Each new paragraph clones the w:pPr (indent, alignment, tabs, spacing) of
    `ref_paragraph`, and each run clones the w:rPr (font family, size, bold, etc.)
    of ref_paragraph's first run. This makes the new paragraphs render with the
    same formatting as the tenant property-address paragraphs that the template
    already aligns correctly.

    When `zero_spacing=True`, the cloned pPr has its w:spacing replaced with an
    explicit before=0/after=0 spacing element so neither the ref's nor the
    BodyText style's default spacing leaks through. Without this, AM saw
    "two blank lines" between 'delivered to:' and the landlord name even with
    the ref's spacing stripped.

    When `left_indent` is provided (twips/dxa), the cloned pPr's w:ind w:left is
    overridden so the inserted paragraphs sit at a specific indent regardless of
    the ref paragraph's own indent. Used to push landlord blocks to the
    signature-line indent instead of body-left.
    """
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    parent = insert_before_element.getparent()
    idx = parent.index(insert_before_element)

    ref_pPr = ref_paragraph._p.find(qn("w:pPr"))
    ref_run = ref_paragraph.runs[0] if ref_paragraph.runs else None
    ref_rPr = ref_run._element.find(qn("w:rPr")) if ref_run is not None else None

    for line in lines:
        p = OxmlElement("w:p")
        if ref_pPr is not None:
            cloned_pPr = deepcopy(ref_pPr)
            if zero_spacing:
                sp = cloned_pPr.find(qn("w:spacing"))
                if sp is not None:
                    cloned_pPr.remove(sp)
                explicit_sp = OxmlElement("w:spacing")
                explicit_sp.set(qn("w:before"), "0")
                explicit_sp.set(qn("w:after"), "0")
                cloned_pPr.append(explicit_sp)
            if left_indent is not None:
                ind = cloned_pPr.find(qn("w:ind"))
                if ind is None:
                    ind = OxmlElement("w:ind")
                    cloned_pPr.append(ind)
                ind.set(qn("w:left"), str(left_indent))
                if ind.get(qn("w:firstLine")):
                    ind.set(qn("w:firstLine"), "0")
            p.append(cloned_pPr)
        r = OxmlElement("w:r")
        if ref_rPr is not None:
            r.append(deepcopy(ref_rPr))
        t = OxmlElement("w:t")
        t.text = line
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        parent.insert(idx, p)
        idx += 1


def _find_body_ref_paragraph(doc):
    """Locate the body paragraph used as the canonical left-indent reference
    for landlord blocks: "Payment of the above..." (w:left=355-393). Both the
    owner block (page 1) and housing-provider block (page 3) clone this
    paragraph's w:pPr so they align with the body text.
    """
    for para in doc.paragraphs:
        if "Payment of the above" in para.text:
            return para
    return None


def _replace_owner_block_with_paragraphs(
    doc, owner_name: str, addr_line1: str, city_state_zip: str
) -> None:
    """Replace the LANDLORD_NAME / COMPANY / ADDRESS template paragraphs with 3 plain
    paragraphs (name / street / city-state-zip), pushed right to the signature-line
    indent (LANDLORD_BLOCK_LEFT_TWIPS).

    Per AM's call: this block must be left-justified TEXT (not centered) but
    indented to the right toward the signature-line position. We clone w:pPr
    from "Payment of the above..." for run/paragraph styling, then override
    w:ind w:left and force explicit zero before/after spacing so the block
    sits as one tight unit directly under "delivered to:" with only one blank
    line between them.
    """
    body = doc.element.body
    paras = list(doc.paragraphs)
    for i, para in enumerate(paras):
        full = "".join(r.text for r in para.runs)
        if "{{LANDLORD_NAME}}" not in full:
            continue
        to_remove = [para._element]
        if i + 1 < len(paras):
            nxt = "".join(r.text for r in paras[i + 1].runs)
            if "{{LANDLORD_COMPANY}}" in nxt or not nxt.strip():
                to_remove.append(paras[i + 1]._element)
        if i + 2 < len(paras):
            nxt2 = "".join(r.text for r in paras[i + 2].runs)
            if "{{LANDLORD_ADDRESS}}" in nxt2:
                to_remove.append(paras[i + 2]._element)

        ref_para = _find_body_ref_paragraph(doc) or (paras[i - 1] if i > 0 else para)

        lines = [owner_name, addr_line1, city_state_zip]
        _insert_paragraphs_before(
            to_remove[0],
            lines,
            ref_para,
            zero_spacing=True,
            left_indent=LANDLORD_BLOCK_LEFT_TWIPS,
        )
        for el in to_remove:
            body.remove(el)
        return


def _compact_empty_paragraphs_before_heading(doc, heading_substring: str) -> int:
    """Remove consecutive empty/whitespace-only paragraphs directly preceding any
    paragraph whose text contains ``heading_substring``.

    On multi-row AMOUNTS_DUE renders (3+ rows) the body lands just past page 2's
    bottom margin, so the "NOTICES FROM THE STATE OF CALIFORNIA:" block gets
    pushed to page 3, leaving page 2 nearly empty. Collapsing the handful of
    throwaway spacer paragraphs the template ships with pulls the heading back
    up so the California block starts on page 2 in those cases.

    Returns the number of empty paragraphs removed.
    """
    from docx.oxml.ns import qn

    body = doc.element.body
    paras = list(doc.paragraphs)
    removed = 0
    for i, para in enumerate(paras):
        if heading_substring not in para.text:
            continue
        j = i - 1
        while j >= 0:
            prev_text = "".join(r.text for r in paras[j].runs).strip()
            if prev_text:
                break
            prev_el = paras[j]._element
            if prev_el.find(qn("w:pPr")) is not None and prev_el.find(qn("w:pPr")).find(qn("w:sectPr")) is not None:
                break
            body.remove(prev_el)
            removed += 1
            j -= 1
        break
    return removed


def _replace_housing_provider_block_with_paragraphs(
    doc, addr_line1: str, city_state_zip: str, phone: str
) -> None:
    """Replace the Housing Provider's / Landlord's Address block with 4 plain paragraphs
    (label / street / city-state-zip / phone) stacked directly beneath the
    "Date: (Original signed by agent/owner)" signature line.

    Reads the signature paragraph's literal w:left and applies the same indent
    to the block so both share the same paragraph indent. Empty paragraphs
    sitting between the signature line and the original block label are
    removed so only one line of gap (w:before=240) remains above the block.
    Internal block lines have w:spacing stripped so the BodyText default
    line-height applies and the four rows stack tight.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    body = doc.element.body
    paras = list(doc.paragraphs)

    sig_para = None
    sig_idx = None
    for idx, p in enumerate(paras):
        if "Original signed by" in p.text:
            sig_para = p
            sig_idx = idx
            break

    # Anchor to the tab-stop position where "(Original signed by agent/owner)"
    # actually renders (the paragraph's w:left=473 only positions "Date:" — the
    # parenthetical lands on the 2nd tab stop). Read tabs dynamically so the
    # alignment tracks any future template tweak; fall back to 3896 (current
    # template's 2nd tab stop) if the structure is missing.
    sig_left = 3896
    if sig_para is not None:
        sig_pPr = sig_para._p.find(qn("w:pPr"))
        if sig_pPr is not None:
            tabs_el = sig_pPr.find(qn("w:tabs"))
            if tabs_el is not None:
                tab_positions = []
                for tab in tabs_el.findall(qn("w:tab")):
                    pos = tab.get(qn("w:pos"))
                    if pos is not None:
                        try:
                            tab_positions.append(int(pos))
                        except ValueError:
                            pass
                if len(tab_positions) >= 2:
                    sig_left = tab_positions[1]

    for i, para in enumerate(paras):
        text = para.text.strip()
        if "Housing Provider" not in text and "Landlord's Address" not in text:
            continue
        label = text
        to_remove = [para._element]
        for j in range(i + 1, min(i + 4, len(paras))):
            nxt = "".join(r.text for r in paras[j].runs).strip()
            if (
                "{{LANDLORD_ADDRESS}}" in nxt
                or "{{LANDLORD_PHONE}}" in nxt
                or "Phone:" in nxt
                or not nxt
            ):
                to_remove.append(paras[j]._element)
            else:
                break

        if sig_idx is not None and sig_idx < i:
            for k in range(sig_idx + 1, i):
                k_text = "".join(r.text for r in paras[k].runs).strip()
                if not k_text:
                    body.remove(paras[k]._element)

        ref_para = _find_body_ref_paragraph(doc) or para
        lines = [label, addr_line1, city_state_zip, f"Phone: {phone}"]
        _insert_paragraphs_before(
            to_remove[0],
            lines,
            ref_para,
            zero_spacing=False,
            left_indent=sig_left,
        )
        for el in to_remove:
            body.remove(el)

        target_texts = [label, addr_line1, city_state_zip, f"Phone: {phone}"]
        found = []
        for p in doc.paragraphs:
            if len(found) >= 4:
                break
            if p.text.strip() == target_texts[len(found)]:
                found.append(p)
        if len(found) == 4:
            pPr0 = _ensure_pPr(found[0])
            sp0 = pPr0.find(qn("w:spacing"))
            if sp0 is None:
                sp0 = OxmlElement("w:spacing")
                pPr0.append(sp0)
            sp0.set(qn("w:before"), "240")
            sp0.set(qn("w:after"), "0")
            for q in found[1:]:
                pPr_q = q._p.find(qn("w:pPr"))
                if pPr_q is not None:
                    sp_q = pPr_q.find(qn("w:spacing"))
                    if sp_q is not None:
                        pPr_q.remove(sp_q)
        return


@app.post("/generate-notice")
async def generate_notice(req: GenerateNoticeRequest):
    """Generate a filled .docx eviction notice from template and return it."""
    logger.info(f"generate-notice: type={req.notice_type} day_count={req.day_count} case={req.case_name}")

    # Validate day count
    if req.day_count not in DAY_COUNT_WORDS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid day_count {req.day_count}. Must be one of: {list(DAY_COUNT_WORDS.keys())}",
        )

    # Select template based on notice_type
    template_name = NOTICE_TEMPLATES.get(req.notice_type)
    if not template_name:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid notice_type '{req.notice_type}'. Must be one of: {list(NOTICE_TEMPLATES.keys())}",
        )
    template_path = TEMPLATES_DIR / template_name
    if not template_path.exists():
        raise HTTPException(status_code=500, detail=f"Template not found: {template_name}")

    # --- Sanitize NOT_FOUND values ---
    BLANK_LINE = "_______________"
    def sanitize(val: str) -> str:
        """Replace NOT_FOUND/empty/None values with a blank line for the document."""
        if not val:
            return BLANK_LINE
        s = str(val).strip()
        if s.upper() in ("NOT_FOUND", "NOT FOUND", "N/A", "NONE", "NULL", ""):
            return BLANK_LINE
        return s

    # Sanitize all request fields before use
    req.tenant_names = sanitize(req.tenant_names)
    # Defense-in-depth against upstream LLM casing drift: enforce
    # "and All Other Occupants" (Title Case) regardless of how the caller
    # punctuated/cased the suffix.
    req.tenant_names = _normalize_all_other_occupants(req.tenant_names)
    req.property_address = sanitize(req.property_address)
    req.property_address_street = sanitize(req.property_address_street) if req.property_address_street else ""
    req.property_address_city = sanitize(req.property_address_city) if req.property_address_city else ""
    # Defense-in-depth: strip any '<City> CA <ZIP>' trailer that leaked into the
    # street field so {{PROPERTY_ADDRESS_STREET}} doesn't duplicate city/state/zip.
    req.property_address_street = _strip_city_state_zip_trailer(req.property_address_street)
    req.county = sanitize(req.county)
    # Defense-in-depth against the n8n lookupCounty/deriveCounty JS functions
    # leaking the bare city-matched value (e.g. "Santa Cruz") instead of the
    # full county name ("Santa Cruz County"). The template renders this as
    # "COUNTY OF {{COUNTY}}", so the suffix matters. Skip if value is the
    # BLANK_LINE sanitization sentinel or already ends in "County".
    if (
        req.county
        and req.county != BLANK_LINE
        and not req.county.strip().lower().endswith("county")
    ):
        req.county = f"{req.county.strip()} County"
    req.total_amount_due = sanitize(req.total_amount_due)
    req.service_date = sanitize(req.service_date)
    req.payment_address = sanitize(req.payment_address)
    req.landlord_name = sanitize(req.landlord_name)
    req.landlord_phone = sanitize(req.landlord_phone)
    req.landlord_address = sanitize(req.landlord_address)
    req.landlord_city = req.landlord_city.strip() if req.landlord_city else ""
    req.landlord_state = req.landlord_state.strip() if req.landlord_state else ""
    req.landlord_zip = req.landlord_zip.strip() if req.landlord_zip else ""
    req.notice_date = sanitize(req.notice_date)
    for a in req.amounts_due:
        a.due_date = sanitize(a.due_date)
        a.amount = sanitize(a.amount)

    # Parse address into street and city/state/zip
    # Prefer explicit street/city fields if provided; otherwise split property_address
    if req.property_address_street:
        street = req.property_address_street
        city_state_zip = req.property_address_city
    else:
        addr_parts = req.property_address.split(",", 1)
        if len(addr_parts) == 2:
            street = addr_parts[0].strip()
            city_state_zip = addr_parts[1].strip()
        else:
            street = req.property_address
            city_state_zip = ""

    # Build landlord address block.
    # Prefer explicit city/state/zip fields (formatted as "City, ST 12345");
    # otherwise parse a combined blob via KNOWN_CITIES so smashed Airtable values
    # like "Santa CruzCA12345" split cleanly into street-line1 + "Santa Cruz, CA 12345".
    landlord_full = req.landlord_address or req.payment_address
    if req.landlord_city or req.landlord_state or req.landlord_zip:
        # When the caller sent structured fields, the 'landlord_full' input represents
        # the street line only; still run it through the parser in case a trailer
        # (e.g. "Main St Santa Cruz CA 12345") leaked in. If the parser finds a
        # trailer, keep the cleaner street portion; otherwise use as-is.
        parsed_line1, _parsed_csz = _parse_landlord_address(landlord_full)
        landlord_addr_line1 = parsed_line1 if parsed_line1 else landlord_full.strip()
        city = req.landlord_city
        state = req.landlord_state
        zipc = req.landlord_zip
        if city and state:
            csz = f"{city}, {state}"
        elif city:
            csz = city
        elif state:
            csz = state
        else:
            csz = ""
        if zipc:
            csz = f"{csz} {zipc}".strip()
        landlord_city_state_zip = csz
    else:
        landlord_addr_line1, landlord_city_state_zip = _parse_landlord_address(landlord_full)

    # Load template
    doc = Document(str(template_path))

    # --- Step 0: Compute NOTICE_HEADER ---
    # Default header based on day_count if not provided
    if req.notice_header:
        notice_header = req.notice_header
    else:
        lower_word, upper_word = DAY_COUNT_WORDS.get(req.day_count, (str(req.day_count), str(req.day_count)))
        notice_header = f"{upper_word} ({req.day_count}) DAY NOTICE TO PAY OR QUIT"

    # --- Step 1: Replace day count in all text (paragraphs + table cells) ---
    for paragraph in doc.paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)
        if "three" in full_text.lower() or "THREE" in full_text or "3 (three)" in full_text:
            _replace_day_count_in_paragraph(paragraph, req.day_count)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = "".join(run.text for run in paragraph.runs)
                    if "three" in full_text.lower() or "THREE" in full_text or "3 (three)" in full_text:
                        _replace_day_count_in_paragraph(paragraph, req.day_count)

    # --- Step 2: Handle the amounts (tab-separated paragraphs or Word table) ---
    amounts_data = [{"due_date": a.due_date, "amount": a.amount} for a in req.amounts_due]
    _fill_amounts_paragraph(doc, amounts_data)

    # --- Step 2b: Replace owner/address blocks with plain paragraphs ---
    _replace_owner_block_with_paragraphs(
        doc, req.landlord_name, landlord_addr_line1, landlord_city_state_zip
    )
    _replace_housing_provider_block_with_paragraphs(
        doc, landlord_addr_line1, landlord_city_state_zip, req.landlord_phone
    )

    # --- Step 2c: Pull California notices up to avoid orphan page 2 ---
    removed = _compact_empty_paragraphs_before_heading(
        doc, "NOTICES FROM THE STATE OF CALIFORNIA"
    )
    if removed:
        logger.info(f"generate-notice: stripped {removed} empty para(s) before California heading")

    # The standalone TOTAL_AMOUNT_DUE Heading1 paragraph is dropped later by
    # `_add_total_row_to_amounts_table` (Step 3e2), which also removes the
    # empty BodyText paragraphs that sit between the rent table and the TOTAL
    # paragraph in the template. Doing both removals there keeps the
    # walk-back logic intact — pre-removing TOTAL here would leave the dead
    # empty paragraphs orphaned above "You are further notified".

    # --- Step 3: Replace standard placeholders ---
    # LANDLORD_NAME, LANDLORD_COMPANY, LANDLORD_ADDRESS, and LANDLORD_PHONE are
    # now handled by the table blocks above, but kept here as fallback in case
    # additional instances exist elsewhere in the template. TOTAL_AMOUNT_DUE
    # is also kept as a no-op fallback for the standalone-paragraph layout.
    fields = {
        "TENANT_NAMES": req.tenant_names,
        "PROPERTY_ADDRESS_STREET": street,
        "PROPERTY_ADDRESS_CITY": city_state_zip,
        "COUNTY": req.county,
        "TOTAL_AMOUNT_DUE": req.total_amount_due,
        "LANDLORD_NAME": req.landlord_name,
        "LANDLORD_COMPANY": "",
        "LANDLORD_ADDRESS": landlord_full,
        "LANDLORD_PHONE": req.landlord_phone,
        # DATE_SERVED is ALWAYS blank — the process server fills it in by hand at
        # service time. n8n has historically sent computed dates here (notice_date,
        # service_date); hardcoding on the Railway side makes rendering immune to
        # whatever upstream workflows end up passing in.
        "DATE_SERVED": "________",
        "NOTICE_HEADER": notice_header,
    }

    # Replace {{KEY}} placeholders in paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, fields)

    # Replace {{KEY}} placeholders in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, fields)

    # Replace in headers/footers
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer is not None:
                for paragraph in header_footer.paragraphs:
                    _replace_in_paragraph(paragraph, fields)

    # --- Step 3b: Tenant-address structural fix (group street/city/COUNTY) ---
    # Run after placeholder substitution so we operate on the rendered text
    # ("123 Tenant Addy" / "Santa Cruz, CA 98598" / "COUNTY OF Santa Cruz County")
    # and can detect landmark paragraphs by content.
    _fix_tenant_address_block(doc)

    # --- Step 3d: Defense-in-depth on "All Other Occupants" casing ---
    # In addition to normalizing req.tenant_names before substitution, also
    # rewrite any post-substitution run text — handles the unlikely case where
    # the LLM injected the phrase elsewhere or the upstream payload bypassed
    # the normalize call.
    _enforce_all_other_occupants_in_doc(doc)

    # --- Step 3e: Force global left alignment on all body + cell paragraphs ---
    # AM: "Everything left-justified. NOTHING centered." Logo header lives in
    # section headers and isn't iterated. The rent table's column-specific
    # alignment (DUE DATE LEFT, AMOUNT DUE RIGHT) is re-applied AFTER this
    # pass via _align_rent_table_columns so it survives.
    _force_left_align_doc(doc)

    # --- Step 3e2: Add the TOTAL AMOUNT DUE row to the rent table ---
    moved_total = _add_total_row_to_amounts_table(doc, req.total_amount_due)
    logger.info(f"generate-notice: total-row inlined into table = {moved_total}")

    # --- Step 3e3: Force rent-table column widths to span body width ---
    # Run AFTER the TOTAL row is added so the new row's two cells also get
    # the body-width column sizing. Sets fixed layout + tblGrid + tcW so the
    # AMOUNT DUE column anchors at the right side of the page (per AM's
    # reference image) and overrides the template's leaked w:jc=center.
    _set_rent_table_column_widths(doc)

    # --- Step 3e4: Apply rent-table column-specific alignment (LAST PASS) ---
    # DUE DATE column → LEFT (dates read from left margin).
    # AMOUNT DUE column → RIGHT (dollar amounts stack on the right edge,
    # standard accounting layout). Must run AFTER _force_left_align_doc
    # (which flattens everything to LEFT) and AFTER _add_total_row so the
    # TOTAL row also gets the column convention.
    _align_rent_table_columns(doc)

    # --- Step 3f: Force California heading onto page 2 for high-row notices ---
    # Below the threshold, the soft compaction already keeps the doc to 2
    # pages without leaving page 1 visibly under-filled. Above it (12-month
    # notices), we force the break so AM gets predictable 2-page output.
    if len(req.amounts_due) >= FORCE_PAGE_BREAK_ROW_THRESHOLD:
        broke = _force_page_break_before(doc, "NOTICES FROM THE STATE OF CALIFORNIA")
        logger.info(
            f"generate-notice: forced page-break-before California for "
            f"{len(req.amounts_due)}-row notice = {broke}"
        )

    # --- Step 4: Save to temp file ---
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc.save(tmp.name)
    file_size = os.path.getsize(tmp.name)
    logger.info(f"generate-notice: saved {tmp.name} size={file_size}")

    # --- Step 5: Use attachments_required from request (deterministic rules in n8n) ---
    attachments = list(req.attachments_required) if req.attachments_required else []

    # --- Step 6: Return the .docx ---
    safe_case_name = req.case_name.replace('"', '').replace("'", "") if req.case_name else "Notice"
    filename = f"{req.day_count}-Day Notice - {safe_case_name}.docx"

    with open(tmp.name, "rb") as f:
        content = f.read()

    # Clean up temp file
    try:
        os.unlink(tmp.name)
    except OSError:
        pass

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"',
    }
    if attachments:
        headers["X-Attachments-Needed"] = ",".join(attachments)

    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
