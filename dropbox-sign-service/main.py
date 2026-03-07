import json
import logging
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

import httpx
from docx import Document
from dropbox_sign import ApiClient, ApiException, Configuration, apis, models
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import PlainTextResponse
from pydantic import BaseModel

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Blueprint Evictions - Dropbox Sign Service")

# --- Config ---

DROPBOX_SIGN_API_KEY = os.environ["DROPBOX_SIGN_API_KEY"]
AIRTABLE_API_KEY = os.environ["AIRTABLE_API_KEY"]
AIRTABLE_BASE_ID = os.environ.get("AIRTABLE_BASE_ID", "appumajkmYLcMryFd")
AIRTABLE_TABLE_ID = os.environ.get("AIRTABLE_TABLE_ID", "tblZbUGz8OTvFNh9i")

configuration = Configuration(username=DROPBOX_SIGN_API_KEY)

TEMPLATES_DIR = Path(__file__).parent / "templates"

TEMPLATE_MAP = {
    "3-Day Pay or Quit": "3day_commercial_TEMPLATE.docx",
    "3-Day Perform or Quit": "3day_perform_quit_TEMPLATE.docx",
    "3-Day Quit": "3day_quit_notice_TEMPLATE.docx",
    "TPO Warning": "tpo_warning_TEMPLATE.docx",
    "TPA Warning": "tpa_warning_TEMPLATE.docx",
}


# --- Models ---


class SendSignatureRequest(BaseModel):
    signer_email: str
    signer_name: str
    document_name: str
    record_id: str
    notice_type: str
    case_name: str
    fields: dict
    file_url: Optional[str] = None


class SendSignatureResponse(BaseModel):
    signature_request_id: str
    status: str
    message: str


# --- Helpers ---


def fill_template(notice_type: str, fields: dict) -> str:
    """Fill a Word template with field values and return the path to the filled .docx."""
    template_name = TEMPLATE_MAP.get(notice_type)
    if not template_name:
        raise ValueError(
            f"Unknown notice_type '{notice_type}'. "
            f"Valid types: {', '.join(TEMPLATE_MAP.keys())}"
        )

    template_path = TEMPLATES_DIR / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    # Replace {{PLACEHOLDER}} tags in paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, fields)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, fields)

    # Replace in headers/footers
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer is not None:
                for paragraph in header_footer.paragraphs:
                    _replace_in_paragraph(paragraph, fields)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc.save(tmp.name)
    return tmp.name


def _replace_in_paragraph(paragraph, fields: dict) -> None:
    """Replace {{KEY}} placeholders in a paragraph while preserving formatting.

    Handles two cases:
    1. Placeholder fits entirely within one run → replace in-place, formatting preserved.
    2. Placeholder spans multiple runs (Word splits text unpredictably) → merge only
       the affected runs into the first one, replace there, clear the rest.
    """
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(run.text for run in runs)
    if "{{" not in full_text:
        return

    # Build list of placeholders to replace
    replacements = {}
    for key, value in fields.items():
        placeholder = "{{" + key + "}}"
        if placeholder in full_text:
            replacements[placeholder] = str(value) if value else ""

    if not replacements:
        return

    # Pass 1: Replace placeholders that fit within a single run
    for run in runs:
        for placeholder, value in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    # Check if any placeholders still remain (they span multiple runs)
    full_text_after = "".join(run.text for run in runs)
    remaining = {p: v for p, v in replacements.items() if p in full_text_after}

    if not remaining:
        return

    # Pass 2: Handle cross-run placeholders by finding the run span
    for placeholder, value in remaining.items():
        while True:
            # Rebuild positions each iteration since runs change
            texts = [run.text for run in runs]
            combined = "".join(texts)
            pos = combined.find(placeholder)
            if pos < 0:
                break

            # Find which runs the placeholder spans
            char_count = 0
            start_run = end_run = None
            for i, t in enumerate(texts):
                run_start = char_count
                run_end = char_count + len(t)
                if start_run is None and pos < run_end:
                    start_run = i
                    offset_in_start = pos - run_start
                if start_run is not None and pos + len(placeholder) <= run_end:
                    end_run = i
                    offset_in_end = pos + len(placeholder) - run_start
                    break
                char_count = run_end

            if start_run is None or end_run is None:
                break

            # Merge the spanned text into the start run, replace, clear others
            merged = "".join(texts[start_run : end_run + 1])
            runs[start_run].text = merged.replace(placeholder, value, 1)
            for j in range(start_run + 1, end_run + 1):
                runs[j].text = ""


def convert_docx_to_pdf(docx_path: str) -> str:
    """Convert a .docx file to PDF using LibreOffice and return the PDF path."""
    output_dir = os.path.dirname(docx_path)
    result = subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            docx_path,
        ],
        capture_output=True,
        text=True,
        timeout=60,
    )

    if result.returncode != 0:
        logger.error(f"LibreOffice conversion failed: {result.stderr}")
        raise RuntimeError(f"PDF conversion failed: {result.stderr}")

    pdf_path = docx_path.rsplit(".", 1)[0] + ".pdf"
    if not os.path.exists(pdf_path):
        raise RuntimeError(
            f"PDF not created at expected path: {pdf_path}. "
            f"LibreOffice stdout: {result.stdout}"
        )

    return pdf_path


def generate_notice_pdf(notice_type: str, fields: dict) -> str:
    """Fill a Word template and convert to PDF. Returns the PDF temp file path."""
    docx_path = fill_template(notice_type, fields)
    try:
        pdf_path = convert_docx_to_pdf(docx_path)
        return pdf_path
    finally:
        try:
            os.unlink(docx_path)
        except OSError:
            pass


async def download_file(url: str) -> str:
    """Download a file from URL to a temp path and return the path."""
    async with httpx.AsyncClient(follow_redirects=True, timeout=60) as client:
        resp = await client.get(url)
        resp.raise_for_status()

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    tmp.write(resp.content)
    tmp.close()
    return tmp.name


async def update_airtable_status(record_id: str, status: str) -> None:
    """Update a record's Status field in Airtable."""
    url = f"https://api.airtable.com/v0/{AIRTABLE_BASE_ID}/{AIRTABLE_TABLE_ID}/{record_id}"
    headers = {
        "Authorization": f"Bearer {AIRTABLE_API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {"fields": {"Status": status}}

    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.patch(url, headers=headers, json=payload)
        if resp.status_code != 200:
            logger.error(
                f"Airtable update failed for {record_id}: {resp.status_code} {resp.text}"
            )
            raise HTTPException(
                status_code=502, detail=f"Airtable update failed: {resp.text}"
            )
        logger.info(f"Airtable record {record_id} updated to Status='{status}'")


# --- Endpoints ---


@app.get("/health")
async def health():
    return {"status": "ok", "service": "dropbox-sign-service"}


@app.post("/send-signature", response_model=SendSignatureResponse)
async def send_signature(req: SendSignatureRequest):
    """Fill a branded Word template (or download from URL) and send to Dropbox Sign."""
    logger.info(
        f"Sending signature request: {req.document_name} to {req.signer_email}"
    )

    # Get PDF: either download from URL or generate from template
    try:
        if req.file_url:
            file_path = await download_file(req.file_url)
        else:
            file_path = generate_notice_pdf(
                notice_type=req.notice_type,
                fields=req.fields,
            )
    except Exception as e:
        logger.error(f"Failed to prepare PDF: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to prepare PDF: {str(e)}")

    try:
        with ApiClient(configuration) as api_client:
            signature_request_api = apis.SignatureRequestApi(api_client)

            signer = models.SubSignatureRequestSigner(
                email_address=req.signer_email,
                name=req.signer_name,
                order=0,
            )

            signing_options = models.SubSigningOptions(
                draw=True,
                type=True,
                upload=True,
                phone=False,
                default_type="draw",
            )

            data = models.SignatureRequestSendRequest(
                title=req.document_name,
                subject=f"Please sign: {req.document_name}",
                message=(
                    f"Hi {req.signer_name},\n\n"
                    f"Please review and sign the attached document: {req.document_name}.\n\n"
                    "Thank you,\nBlueprint Evictions LLC"
                ),
                signers=[signer],
                files=[open(file_path, "rb")],
                metadata={"record_id": req.record_id},
                signing_options=signing_options,
                test_mode=os.environ.get("DROPBOX_SIGN_TEST_MODE", "0") == "1",
            )

            result = signature_request_api.signature_request_send(data)
            sig_request = result.signature_request

            logger.info(
                f"Signature request created: {sig_request.signature_request_id}"
            )

            return SendSignatureResponse(
                signature_request_id=sig_request.signature_request_id,
                status="sent",
                message=f"Signature request sent to {req.signer_email}",
            )

    except ApiException as e:
        logger.error(f"Dropbox Sign API error: {e.status} {e.body}")
        raise HTTPException(
            status_code=502,
            detail=f"Dropbox Sign API error: {e.body}",
        )
    finally:
        # Clean up temp file
        try:
            os.unlink(file_path)
        except OSError:
            pass


@app.post("/signature-callback")
async def signature_callback(request: Request):
    """Webhook endpoint that Dropbox Sign calls when a document is signed.

    Dropbox Sign sends event data as form-encoded with a 'json' field.
    On event 'signature_request_all_signed', update Airtable Status to 'Signed (A)'.
    """
    form = await request.form()

    # Dropbox Sign sends the payload in a 'json' form field
    json_str = form.get("json")
    if not json_str:
        logger.warning("Callback received with no 'json' field")
        return {"status": "ignored", "reason": "no json payload"}

    payload = json.loads(json_str)
    event = payload.get("event", {})
    event_type = event.get("event_type")
    event_time = event.get("event_time")

    logger.info(f"Dropbox Sign callback: event_type={event_type} time={event_time}")

    # Respond to the callback test (Dropbox Sign sends this to verify the endpoint)
    if event_type == "callback_test":
        logger.info("Responding to callback_test with 'Hello API Event Received'")
        return PlainTextResponse("Hello API Event Received")

    # Only act on all-signed events
    if event_type != "signature_request_all_signed":
        logger.info(f"Ignoring event type: {event_type}")
        return {"status": "ignored", "event_type": event_type}

    # Extract metadata with record_id
    signature_request = payload.get("signature_request", {})
    metadata = signature_request.get("metadata", {})
    record_id = metadata.get("record_id")
    sig_request_id = signature_request.get("signature_request_id", "unknown")

    if not record_id:
        logger.error(
            f"No record_id in metadata for signature request {sig_request_id}"
        )
        return {"status": "error", "reason": "no record_id in metadata"}

    logger.info(
        f"All signed for request {sig_request_id}, updating record {record_id}"
    )

    # Update Airtable
    await update_airtable_status(record_id, "Signed (A)")

    return {
        "status": "processed",
        "event_type": event_type,
        "record_id": record_id,
        "signature_request_id": sig_request_id,
    }
